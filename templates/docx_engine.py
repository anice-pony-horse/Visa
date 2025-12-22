"""
DOCX Template Engine - Professional Legal Document Generation
==============================================================

Generates law-firm quality DOCX documents for visa petitions.

CRITICAL REQUIREMENTS:
1. Output CLEAN PROSE - no markdown artifacts
2. Professional formatting (Times New Roman 12pt, 1" margins)
3. Proper legal citations (8 C.F.R. format)
4. USCIS-compliant structure

Document Types:
- Cover Letter: Brief intro letter to USCIS
- Legal Brief: Full petition letter with criterion analysis
- Table of Contents: Professional TOC with page references
- Comparable Evidence Letter: CE explanation for O-1A/O-1B/EB-1A
- Filing Instructions: DIY guide for self-petitioners
"""

import os
from datetime import datetime
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Check for python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Install with: pip install python-docx")


@dataclass
class CaseData:
    """Case data for document generation"""
    beneficiary_name: str
    petitioner_name: str
    visa_type: str
    nationality: str = ""
    field: str = ""
    job_title: str = ""
    duration: str = "3 years"
    processing_type: str = "Regular"
    service_center: str = "California Service Center"
    filing_fee: str = "$460"
    premium_fee: str = "$2,805"
    criteria_met: List[str] = None

    def __post_init__(self):
        if self.criteria_met is None:
            self.criteria_met = []


# CFR citations by visa type
CFR_CITATIONS = {
    "O-1A": {
        "main": "8 C.F.R. § 214.2(o)(3)(iii)",
        "standard": "INA § 101(a)(15)(O)(i); 8 U.S.C. § 1101(a)(15)(O)(i)",
        "comparable": "8 C.F.R. § 214.2(o)(3)(v)",
        "criteria": {
            "A": "8 C.F.R. § 214.2(o)(3)(iii)(A) - Awards and Prizes",
            "B": "8 C.F.R. § 214.2(o)(3)(iii)(B) - Membership",
            "C": "8 C.F.R. § 214.2(o)(3)(iii)(C) - Published Material",
            "D": "8 C.F.R. § 214.2(o)(3)(iii)(D) - Judging",
            "E": "8 C.F.R. § 214.2(o)(3)(iii)(E) - Original Contributions",
            "F": "8 C.F.R. § 214.2(o)(3)(iii)(F) - Authorship",
            "G": "8 C.F.R. § 214.2(o)(3)(iii)(G) - Critical Employment",
            "H": "8 C.F.R. § 214.2(o)(3)(iii)(H) - High Remuneration"
        }
    },
    "P-1A": {
        "main": "8 C.F.R. § 214.2(p)(4)(ii)(B)",
        "standard": "INA § 101(a)(15)(P)(i); 8 U.S.C. § 1101(a)(15)(P)(i)",
        "comparable": None,  # P-1A has NO comparable evidence
        "criteria": {
            "A": "8 C.F.R. § 214.2(p)(4)(ii)(B)(1) - International Recognition",
            "B": "8 C.F.R. § 214.2(p)(4)(ii)(B)(2) - Team Achievements",
            "C": "8 C.F.R. § 214.2(p)(4)(ii)(B)(3) - Awards",
            "D": "8 C.F.R. § 214.2(p)(4)(ii)(B)(4) - Ranking",
            "E": "8 C.F.R. § 214.2(p)(4)(ii)(B)(5) - Media Coverage"
        }
    },
    "EB-1A": {
        "main": "8 C.F.R. § 204.5(h)(3)",
        "standard": "INA § 203(b)(1)(A); 8 U.S.C. § 1153(b)(1)(A)",
        "comparable": "8 C.F.R. § 204.5(h)(4)",
        "criteria": {
            "A": "8 C.F.R. § 204.5(h)(3)(i) - Awards",
            "B": "8 C.F.R. § 204.5(h)(3)(ii) - Membership",
            "C": "8 C.F.R. § 204.5(h)(3)(iii) - Published Material",
            "D": "8 C.F.R. § 204.5(h)(3)(iv) - Judging",
            "E": "8 C.F.R. § 204.5(h)(3)(v) - Original Contributions",
            "F": "8 C.F.R. § 204.5(h)(3)(vi) - Authorship",
            "G": "8 C.F.R. § 204.5(h)(3)(vii) - Artistic Exhibitions",
            "H": "8 C.F.R. § 204.5(h)(3)(viii) - Leading Role",
            "I": "8 C.F.R. § 204.5(h)(3)(ix) - High Salary",
            "J": "8 C.F.R. § 204.5(h)(3)(x) - Commercial Success"
        }
    }
}


class DOCXTemplateEngine:
    """
    Professional DOCX document generator for visa petitions.

    Produces law-firm quality output with:
    - Times New Roman 12pt font
    - 1-inch margins
    - Proper legal formatting
    - Clean prose (NO markdown)
    """

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required. Install with: pip install python-docx")

    def _create_document(self) -> Document:
        """Create base document with professional styling."""
        doc = Document()

        # Set margins (1 inch)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        return doc

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Add a heading without markdown symbols."""
        heading = doc.add_heading(text.upper(), level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14 if level == 1 else 12)

    def _add_paragraph(self, doc: Document, text: str, indent: bool = True, bold: bool = False):
        """Add a paragraph with proper formatting."""
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold

    def _add_signature_block(self, doc: Document, name: str = "", date: bool = True):
        """Add professional signature block."""
        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Respectfully submitted,")

        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("_" * 30)

        if name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(name)

        if date:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(datetime.now().strftime("%B %d, %Y"))

    def generate_cover_letter(
        self,
        case: CaseData,
        exhibits: List[Dict[str, Any]],
        output_path: str
    ) -> str:
        """
        Generate professional cover letter.

        Args:
            case: Case data
            exhibits: List of exhibits
            output_path: Output file path

        Returns:
            Path to generated document
        """
        doc = self._create_document()

        # Date (right aligned)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(datetime.now().strftime("%B %d, %Y"))

        doc.add_paragraph()

        # Address block
        address_lines = [
            "U.S. Citizenship and Immigration Services",
            case.service_center,
            ""  # Address would go here
        ]
        for line in address_lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(0)

        doc.add_paragraph()

        # Re: line
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"Re:    Form I-129 Petition for {case.visa_type} Nonimmigrant Status")
        run.bold = True

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.75)
        p.add_run(f"Beneficiary: {case.beneficiary_name}")

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.75)
        p.add_run(f"Petitioner: {case.petitioner_name}")

        doc.add_paragraph()

        # Salutation
        doc.add_paragraph("Dear Sir or Madam:")

        doc.add_paragraph()

        # Body paragraphs
        intro = (
            f"Please accept this petition filed on behalf of {case.beneficiary_name} "
            f'("Beneficiary" or "{case.beneficiary_name.split()[-1]}"), '
            f"a citizen of {case.nationality or '[Country]'}, for classification "
            f"as a nonimmigrant worker under {CFR_CITATIONS.get(case.visa_type, {}).get('standard', 'the Immigration and Nationality Act')}."
        )
        self._add_paragraph(doc, intro)

        petitioner_para = (
            f'{case.petitioner_name} ("Petitioner") seeks to employ the Beneficiary as '
            f"{case.job_title or '[Job Title]'} for a period of {case.duration}. "
            f"The Beneficiary possesses extraordinary ability in {case.field or '[field]'}, "
            f"as demonstrated by the enclosed evidence satisfying {len(case.criteria_met) or 3} "
            f"of the regulatory criteria set forth at {CFR_CITATIONS.get(case.visa_type, {}).get('main', '8 C.F.R.')}."
        )
        self._add_paragraph(doc, petitioner_para)

        # Exhibit list intro
        self._add_paragraph(doc, "This petition is accompanied by the following:")

        doc.add_paragraph()

        # Forms list
        forms = [
            "Form I-129 with O and P Classifications Supplement",
            "Form G-1450 (Credit Card Authorization)",
        ]
        if case.processing_type == "Premium":
            forms.append("Form I-907 (Premium Processing Request)")
        forms.extend([
            f"Filing fee of {case.filing_fee}",
        ])
        if case.processing_type == "Premium":
            forms.append(f"Premium processing fee of {case.premium_fee}")
        forms.append(f"Supporting documentation organized as Exhibits A through {chr(64 + len(exhibits))}")

        for i, form in enumerate(forms):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(f"{i + 1}.  {form}")

        doc.add_paragraph()

        # Exhibit list
        self._add_paragraph(doc, "The following exhibits are submitted in support of this petition:")

        doc.add_paragraph()

        for exhibit in exhibits[:15]:  # Limit to first 15 in cover letter
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            num = exhibit.get("exhibit_number", exhibit.get("number", ""))
            name = exhibit.get("name", exhibit.get("filename", "Document"))
            p.add_run(f"Exhibit {num}:  {name}")

        if len(exhibits) > 15:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(f"[... and {len(exhibits) - 15} additional exhibits]")

        doc.add_paragraph()

        # Closing
        self._add_paragraph(doc, "We respectfully request approval of this petition.")

        # Signature
        self._add_signature_block(doc)

        # Save
        doc.save(output_path)
        logger.info(f"Generated cover letter: {output_path}")
        return output_path

    def generate_legal_brief(
        self,
        case: CaseData,
        exhibits: List[Dict[str, Any]],
        criterion_analyses: Dict[str, str],
        output_path: str
    ) -> str:
        """
        Generate full legal brief / petition letter.

        Args:
            case: Case data
            exhibits: List of exhibits
            criterion_analyses: Dict mapping criterion letter to analysis text
            output_path: Output file path

        Returns:
            Path to generated document
        """
        doc = self._create_document()

        # Title
        self._add_heading(doc, "PETITION LETTER", 1)
        self._add_heading(doc, "IN SUPPORT OF I-129 PETITION", 2)
        self._add_heading(doc, f"FOR {case.visa_type} NONIMMIGRANT CLASSIFICATION", 2)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Beneficiary: {case.beneficiary_name}")

        doc.add_paragraph()

        # I. INTRODUCTION
        self._add_heading(doc, "I. INTRODUCTION", 2)

        intro = (
            f"This petition is filed on behalf of {case.beneficiary_name}, a citizen of "
            f"{case.nationality or '[Country]'}, who seeks {case.visa_type} classification "
            f"based on extraordinary ability in {case.field or '[field]'}. The Beneficiary "
            f"satisfies the regulatory criteria through evidence of outstanding achievements "
            f"as detailed below."
        )
        self._add_paragraph(doc, intro)

        # II. STATEMENT OF FACTS
        self._add_heading(doc, "II. STATEMENT OF FACTS", 2)

        p = doc.add_paragraph()
        run = p.add_run("A. Beneficiary's Background")
        run.bold = True

        background = (
            f"{case.beneficiary_name} is a distinguished professional in the field of "
            f"{case.field or '[field]'}. The Beneficiary has achieved recognition through "
            f"numerous accomplishments as documented in the enclosed exhibits."
        )
        self._add_paragraph(doc, background)

        # III. LEGAL STANDARD
        self._add_heading(doc, "III. LEGAL STANDARD", 2)

        citations = CFR_CITATIONS.get(case.visa_type, CFR_CITATIONS["O-1A"])

        legal_standard = (
            f"The {case.visa_type} classification is available to individuals who possess "
            f"extraordinary ability in their field of endeavor. See {citations['standard']}."
        )
        self._add_paragraph(doc, legal_standard)

        extraordinary_def = (
            '"Extraordinary ability" is defined as "a level of expertise indicating that '
            'the individual is one of that small percentage who have risen to the very top '
            f'of the field of endeavor." {citations["main"]}.'
        )
        self._add_paragraph(doc, extraordinary_def)

        criteria_intro = (
            "To establish eligibility, the beneficiary must demonstrate sustained national "
            f"or international acclaim by satisfying at least three of the criteria set forth "
            f"at {citations['main']}, or by providing evidence of a major, internationally "
            "recognized award."
        )
        self._add_paragraph(doc, criteria_intro)

        # IV. CRITERION-BY-CRITERION ANALYSIS
        self._add_heading(doc, "IV. CRITERION-BY-CRITERION ANALYSIS", 2)

        claimed_count = len([c for c in case.criteria_met if c])
        self._add_paragraph(
            doc,
            f"The Beneficiary satisfies {claimed_count} of the regulatory criteria:"
        )

        # Add each criterion analysis
        for letter in sorted(criterion_analyses.keys()):
            criterion_text = citations.get("criteria", {}).get(letter, f"Criterion {letter}")

            p = doc.add_paragraph()
            run = p.add_run(f"A. CRITERION ({letter}): {criterion_text.split(' - ')[-1].upper()}")
            run.bold = True

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(criterion_text)

            # Regulatory Requirement
            p = doc.add_paragraph()
            run = p.add_run("Regulatory Requirement:")
            run.bold = True
            run.underline = True

            # Analysis text (cleaned of any markdown)
            analysis = criterion_analyses[letter]
            analysis = analysis.replace("**", "").replace("##", "").replace("*", "")
            self._add_paragraph(doc, analysis)

            doc.add_paragraph()

        # V. CONCLUSION
        self._add_heading(doc, "V. CONCLUSION", 2)

        conclusion = (
            f"Based on the foregoing, the Beneficiary has demonstrated extraordinary ability "
            f"in {case.field or '[field]'} by satisfying {claimed_count} of the regulatory "
            f"criteria set forth at {citations['main']}. The totality of the evidence "
            "establishes that the Beneficiary has risen to the very top of the field and "
            "has achieved sustained national and international acclaim."
        )
        self._add_paragraph(doc, conclusion)

        self._add_paragraph(doc, "We respectfully request that this petition be approved.")

        # Signature
        self._add_signature_block(doc)

        # Save
        doc.save(output_path)
        logger.info(f"Generated legal brief: {output_path}")
        return output_path

    def generate_toc(
        self,
        exhibits: List[Dict[str, Any]],
        case: CaseData,
        output_path: str
    ) -> str:
        """
        Generate professional Table of Contents.

        Args:
            exhibits: List of exhibits with name, page_count, criterion
            case: Case data
            output_path: Output file path

        Returns:
            Path to generated document
        """
        doc = self._create_document()

        # Title
        self._add_heading(doc, "TABLE OF CONTENTS", 1)
        self._add_heading(doc, f"{case.visa_type} PETITION", 2)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Beneficiary: {case.beneficiary_name}")

        doc.add_paragraph()
        doc.add_paragraph()

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Exhibit"
        header_cells[1].text = "Description"
        header_cells[2].text = "Pages"

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add exhibits
        current_section = ""
        page_num = 1

        for exhibit in exhibits:
            # Check for section change
            criterion = exhibit.get("criterion_letter", "")
            if criterion and criterion != current_section:
                current_section = criterion
                # Add section header row
                row = table.add_row()
                row.cells[0].merge(row.cells[2])
                p = row.cells[0].paragraphs[0]
                run = p.add_run(f"CRITERION {criterion}")
                run.bold = True

            # Add exhibit row
            row = table.add_row()
            row.cells[0].text = f"Exhibit {exhibit.get('exhibit_number', exhibit.get('number', ''))}"
            row.cells[1].text = exhibit.get("name", exhibit.get("filename", "Document"))[:60]
            pages = exhibit.get("page_count", exhibit.get("pages", 1))
            row.cells[2].text = str(pages)

            page_num += pages

        doc.add_paragraph()

        # Total
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"TOTAL PAGES: {page_num}")
        run.bold = True

        # Save
        doc.save(output_path)
        logger.info(f"Generated TOC: {output_path}")
        return output_path

    def generate_ce_letter(
        self,
        case: CaseData,
        criterion: str,
        reason_not_applicable: str,
        comparable_evidence: str,
        output_path: str
    ) -> str:
        """
        Generate Comparable Evidence explanation letter.

        IMPORTANT: P-1A has NO comparable evidence provision.
        This should only be called for O-1A, O-1B, or EB-1A.

        Args:
            case: Case data
            criterion: Criterion letter (A, B, C, etc.)
            reason_not_applicable: Why standard criterion doesn't apply
            comparable_evidence: Description of comparable evidence
            output_path: Output file path

        Returns:
            Path to generated document
        """
        # CRITICAL: P-1A has NO comparable evidence
        if case.visa_type == "P-1A":
            raise ValueError("P-1A visa has NO comparable evidence provision. Cannot generate CE letter.")

        doc = self._create_document()

        citations = CFR_CITATIONS.get(case.visa_type, CFR_CITATIONS["O-1A"])
        ce_citation = citations.get("comparable")

        if not ce_citation:
            raise ValueError(f"{case.visa_type} does not support comparable evidence")

        # Title
        self._add_heading(doc, "COMPARABLE EVIDENCE EXPLANATION", 1)
        self._add_heading(doc, f"CRITERION {criterion}", 2)

        doc.add_paragraph()

        # Intro
        intro = (
            f"Pursuant to {ce_citation}, the Petitioner submits comparable evidence "
            f"to establish the Beneficiary's eligibility under Criterion {criterion} "
            f"of the {case.visa_type} regulatory framework."
        )
        self._add_paragraph(doc, intro)

        doc.add_paragraph()

        # Section I
        p = doc.add_paragraph()
        run = p.add_run("I. WHY THE STANDARD CRITERION DOES NOT READILY APPLY")
        run.bold = True

        standard_criterion = citations.get("criteria", {}).get(criterion, f"Criterion {criterion}")

        self._add_paragraph(
            doc,
            f'The standard criterion at {standard_criterion.split(" - ")[0]} requires evidence of '
            f'"{standard_criterion.split(" - ")[-1] if " - " in standard_criterion else "the specified requirement"}."'
        )

        self._add_paragraph(doc, "This criterion does not readily apply to the Beneficiary's occupation for the following reasons:")

        # Clean the reason text of markdown
        reason_clean = reason_not_applicable.replace("**", "").replace("##", "").replace("*", "")
        self._add_paragraph(doc, reason_clean)

        doc.add_paragraph()

        # Section II
        p = doc.add_paragraph()
        run = p.add_run("II. COMPARABLE EVIDENCE SUBMITTED")
        run.bold = True

        self._add_paragraph(doc, "In lieu of the standard evidence, the Petitioner submits:")

        # Clean the evidence text
        evidence_clean = comparable_evidence.replace("**", "").replace("##", "").replace("*", "")
        self._add_paragraph(doc, evidence_clean)

        doc.add_paragraph()

        # Section III
        p = doc.add_paragraph()
        run = p.add_run("III. ANALYSIS OF COMPARABLE SIGNIFICANCE")
        run.bold = True

        analysis = (
            "The submitted evidence is of comparable significance to the standard criterion "
            "because it demonstrates the same level of achievement and recognition that the "
            "standard criterion was designed to identify. The evidence shows that the Beneficiary "
            "has attained a level of expertise indicating they are among the small percentage "
            "who have risen to the very top of their field."
        )
        self._add_paragraph(doc, analysis)

        doc.add_paragraph()

        # Closing
        closing = (
            f"For the foregoing reasons, the Petitioner has established that the Beneficiary "
            f"meets Criterion {criterion} through comparable evidence as permitted by {ce_citation}."
        )
        self._add_paragraph(doc, closing)

        # Save
        doc.save(output_path)
        logger.info(f"Generated CE letter: {output_path}")
        return output_path

    def generate_filing_instructions(
        self,
        case: CaseData,
        exhibits: List[Dict[str, Any]],
        output_path: str
    ) -> str:
        """
        Generate filing instructions for DIY self-petitioners.

        Args:
            case: Case data
            exhibits: List of exhibits
            output_path: Output file path

        Returns:
            Path to generated document
        """
        doc = self._create_document()

        # Title
        self._add_heading(doc, "FILING INSTRUCTIONS", 1)
        self._add_heading(doc, f"{case.visa_type} Self-Petition Package", 2)

        doc.add_paragraph()

        # Before You File
        p = doc.add_paragraph()
        run = p.add_run("BEFORE YOU FILE")
        run.bold = True
        run.underline = True

        checklist = [
            "Review all documents for completeness",
            "Make a complete copy for your records",
            "Ensure all forms are signed in ink (blue or black)",
            "Calculate correct fees",
            "Verify service center address is current"
        ]

        for item in checklist:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(f"[ ] {item}")

        doc.add_paragraph()

        # Filing Checklist
        p = doc.add_paragraph()
        run = p.add_run("FILING CHECKLIST")
        run.bold = True
        run.underline = True

        forms = [
            "Form I-129 (signed and dated)",
            "O/P Supplement (signed)",
            "Form G-1450 (if paying by credit card)",
        ]
        if case.processing_type == "Premium":
            forms.append("Form I-907 (for premium processing)")

        forms.extend([
            f"Filing fee: {case.filing_fee}",
        ])
        if case.processing_type == "Premium":
            forms.append(f"Premium processing fee: {case.premium_fee}")

        forms.append(f"All exhibits in order A through {chr(64 + len(exhibits))}")

        for item in forms:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(f"[ ] {item}")

        doc.add_paragraph()

        # Where to File
        p = doc.add_paragraph()
        run = p.add_run("WHERE TO FILE")
        run.bold = True
        run.underline = True

        self._add_paragraph(doc, "Send your complete package to:", indent=False)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(case.service_center)

        doc.add_paragraph()

        # What to Expect
        p = doc.add_paragraph()
        run = p.add_run("WHAT TO EXPECT")
        run.bold = True
        run.underline = True

        expectations = [
            ("Receipt Notice", "2-4 weeks after filing"),
            ("Premium Processing", "Decision within 15 business days" if case.processing_type == "Premium" else "Not selected"),
            ("Regular Processing", "3-6 months (varies)")
        ]

        for label, detail in expectations:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(detail)

        doc.add_paragraph()

        # RFE Instructions
        p = doc.add_paragraph()
        run = p.add_run("IF YOU RECEIVE A REQUEST FOR EVIDENCE (RFE)")
        run.bold = True
        run.underline = True

        rfe_items = [
            "Respond by the deadline (usually 84 days)",
            "Address EVERY point raised in the RFE",
            "Submit additional evidence as needed",
            "Keep a copy of your response",
            "Use certified mail with tracking"
        ]

        for item in rfe_items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(f"- {item}")

        doc.add_paragraph()

        # Contact Info
        p = doc.add_paragraph()
        run = p.add_run("CONTACT INFORMATION")
        run.bold = True
        run.underline = True

        contacts = [
            ("USCIS Customer Service", "1-800-375-5283"),
            ("Case Status Online", "uscis.gov/casestatus"),
            ("USCIS Website", "uscis.gov")
        ]

        for label, detail in contacts:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(detail)

        # Save
        doc.save(output_path)
        logger.info(f"Generated filing instructions: {output_path}")
        return output_path


# Convenience functions
def generate_cover_letter(case_data: Dict, exhibits: List, output_path: str) -> str:
    """Generate cover letter (convenience function)."""
    engine = DOCXTemplateEngine()
    case = CaseData(**case_data)
    return engine.generate_cover_letter(case, exhibits, output_path)


def generate_legal_brief(case_data: Dict, exhibits: List, analyses: Dict, output_path: str) -> str:
    """Generate legal brief (convenience function)."""
    engine = DOCXTemplateEngine()
    case = CaseData(**case_data)
    return engine.generate_legal_brief(case, exhibits, analyses, output_path)


def generate_toc(exhibits: List, case_data: Dict, output_path: str) -> str:
    """Generate table of contents (convenience function)."""
    engine = DOCXTemplateEngine()
    case = CaseData(**case_data)
    return engine.generate_toc(exhibits, case, output_path)


def generate_ce_letter(case_data: Dict, criterion: str, reason: str, evidence: str, output_path: str) -> str:
    """Generate comparable evidence letter (convenience function)."""
    engine = DOCXTemplateEngine()
    case = CaseData(**case_data)
    return engine.generate_ce_letter(case, criterion, reason, evidence, output_path)


def generate_filing_instructions(case_data: Dict, exhibits: List, output_path: str) -> str:
    """Generate filing instructions (convenience function)."""
    engine = DOCXTemplateEngine()
    case = CaseData(**case_data)
    return engine.generate_filing_instructions(case, exhibits, output_path)
